import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Set

logger = logging.getLogger(__name__)

# Tier 3 documents to exclude from factual RAG indexing according to Phase 4A
TIER_3_EXCLUDED_FILES: Set[str] = {
    "multi-intent.md",
    "multi-intent.txt",
    "follow-up-contextual.md",
    "follow-up-contextual.txt",
    "ambiguous-incomplete.md",
    "ambiguous-incomplete.txt",
}

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    """Represents a raw loaded document with its source metadata."""
    document_id: str
    filename: str
    file_path: Path
    raw_content: str
    file_type: str
    tier: int


class DocumentLoader:
    """
    Loads knowledge documents (.md, .txt, .pdf, .docx) from a specified directory,
    filtering eligible files according to the PERC Phase 4A classification.
    """
    def __init__(self, directory_path: Path):
        self.directory_path = Path(directory_path)

    def get_tier(self, filename: str) -> int:
        """Determines the document tier based on Phase 4A taxonomy."""
        stem = Path(filename).stem.lower()
        if filename in TIER_3_EXCLUDED_FILES or stem in {"multi-intent", "follow-up-contextual", "ambiguous-incomplete"}:
            return 3
        tier_1_stems = {
            "course-discovery",
            "course-details",
            "fees-pricing",
            "eligibility",
            "branch-location",
            "admission-process",
            "availability-status",
        }
        if stem in tier_1_stems:
            return 1
        return 2  # Tier 2 (Pure Unstructured Knowledge)

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extracts text content from a PDF file using pypdf."""
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            extracted_pages = []
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    extracted_pages.append(f"## Page {i}\n\n{text.strip()}")
            return "\n\n".join(extracted_pages)
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return ""

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extracts text content from a Word document (.docx) using python-docx."""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = []
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                if p.style and p.style.name and p.style.name.startswith("Heading 1"):
                    paragraphs.append(f"# {text}")
                elif p.style and p.style.name and p.style.name.startswith("Heading 2"):
                    paragraphs.append(f"## {text}")
                elif p.style and p.style.name and p.style.name.startswith("Heading 3"):
                    paragraphs.append(f"### {text}")
                else:
                    paragraphs.append(text)
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return ""

    def load_document(self, file_path: Path) -> LoadedDocument:
        """Reads a single document file preserving text formatting."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {path}")

        filename = path.name
        doc_id = path.stem
        ext = path.suffix.lower()

        if ext in {".md", ".txt"}:
            raw_content = path.read_text(encoding="utf-8-sig")
        elif ext == ".pdf":
            raw_content = self._extract_text_from_pdf(path)
        elif ext == ".docx":
            raw_content = self._extract_text_from_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        tier = self.get_tier(filename)

        return LoadedDocument(
            document_id=doc_id,
            filename=filename,
            file_path=path,
            raw_content=raw_content,
            file_type=ext.lstrip("."),
            tier=tier,
        )

    def discover_eligible_documents(self, include_tier_3: bool = False) -> List[LoadedDocument]:
        """
        Discovers and loads all eligible documents (.md, .txt, .pdf, .docx) in the directory.
        By default, Tier 3 evaluation/scenario files are strictly excluded.
        """
        if not self.directory_path.exists() or not self.directory_path.is_dir():
            raise NotADirectoryError(f"Directory does not exist: {self.directory_path}")

        all_files = []
        for ext in SUPPORTED_EXTENSIONS:
            all_files.extend(self.directory_path.glob(f"*{ext}"))

        all_files = sorted(all_files, key=lambda p: p.name.lower())
        loaded_docs: List[LoadedDocument] = []

        for f in all_files:
            tier = self.get_tier(f.name)
            if tier == 3 and not include_tier_3:
                continue
            loaded = self.load_document(f)
            if loaded.raw_content.strip():
                loaded_docs.append(loaded)

        return loaded_docs

