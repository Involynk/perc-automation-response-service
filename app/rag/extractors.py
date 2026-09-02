import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Optional


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx", ".html", ".htm", ".csv"}


@dataclass
class ExtractedDocument:
    filename: str
    content_type: str
    text: str


class UnsupportedDocumentTypeError(ValueError):
    """Raised when an uploaded file type cannot be converted to text."""


def _content_type_for(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    return {
        ".md": "text/markdown",
        ".txt": "text/plain",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".html": "text/html",
        ".htm": "text/html",
        ".csv": "text/csv",
    }.get(ext, "application/octet-stream")


def _decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _strip_html(html: str) -> str:
    text = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", html)
    text = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    return re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", text)).strip()


def _extract_pdf(data: bytes, filename: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise UnsupportedDocumentTypeError(
            "PDF support requires pypdf. Install it with `pip install pypdf`."
        ) from exc

    reader = PdfReader(BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages.append(f"## Page {index}\n\n{page_text}")
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError(f"No extractable text found in PDF: {filename}")
    title = Path(filename).stem.replace("-", " ").replace("_", " ").title()
    return f"# {title}\n\n{text}"


def _extract_docx(data: bytes, filename: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise UnsupportedDocumentTypeError(
            "DOCX support requires python-docx. Install it with `pip install python-docx`."
        ) from exc

    document = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if not line:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if "heading 1" in style:
            parts.append(f"# {line}")
        elif "heading 2" in style:
            parts.append(f"## {line}")
        elif "heading 3" in style:
            parts.append(f"### {line}")
        else:
            parts.append(line)
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError(f"No extractable text found in DOCX: {filename}")
    return text


def extract_text(filename: str, data: bytes, content_type: Optional[str] = None) -> ExtractedDocument:
    """Convert an uploaded file into UTF-8 text suitable for the markdown chunker."""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentTypeError(
            f"Unsupported file type '{ext or filename}'. "
            f"Allowed: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    resolved_type = content_type or _content_type_for(filename)

    if ext == ".pdf":
        text = _extract_pdf(data, filename)
    elif ext == ".docx":
        text = _extract_docx(data, filename)
    elif ext in {".html", ".htm"}:
        text = _strip_html(_decode_bytes(data))
        if not text:
            raise ValueError(f"No extractable text found in HTML: {filename}")
        title = Path(filename).stem.replace("-", " ").replace("_", " ").title()
        text = f"# {title}\n\n{text}"
    else:
        text = _decode_bytes(data).strip()
        if not text:
            raise ValueError(f"File is empty: {filename}")

    return ExtractedDocument(filename=filename, content_type=resolved_type, text=text)
